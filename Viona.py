import customtkinter as ctk
from CTkMessagebox import CTkMessagebox
from tkinter.filedialog import askdirectory
from PIL import Image
from pathlib import Path
import glob
import os
from docxtpl import DocxTemplate
from docx import Document

import shutil

TextColor = '#c2c4c3'
BackgroundColor = '#3E3D40'
ButtonColor = '#717173'
HoverColor = '#b8b8b9'
BorderColor = '#aaacaa'

class HandleSubmit():
    
    # --- Check if the user wants a new outputfolder, if yes, create one and change value in dictionary ---
    def New_dir(Dict) -> dict:   
        if Dict['NeueOrdner'] == 1:
            outputdir = f"{Dict['Teilnehmerordner']}{os.sep}{Dict['Name']}, {Dict['Vorname']}"
            Path(outputdir).mkdir(parents=True, exist_ok=True)
            Dict['Teilnehmerordner'] = outputdir
        return Dict

    # --- Test if path exists and return bool for input and output seperatly ---
    def Test_path(dict) -> bool:

        # --- if either path does not exists, show an error and return False ---
        if Path(dict['Vorlagen']).exists() == False or Path(dict['Teilnehmerordner']).exists() == False:
            Popup(Title='Problem mit Ordner!',
                Message=f'Vorlagenordner existiert: {Path(dict["Vorlagen"]).exists()} \n Teilnehmerordner existiert: {Path(dict["Teilnehmerordner"]).exists()}',
                Icon='cancel')
            return False
        else:
            # --- If both paths exists return True ---
            return True
    
    # --- Handle missing values - which could not be provided yet - in dictionary ---
    def Clean_dict(Dict) -> dict:

        # Key for missing values, used for formatting later on
        NoInput = 'Information fehlt'
        for key in list(Dict):
            if Dict[key] == '-':
                Dict[key] = NoInput
            elif Dict[key] == 'Vollzeit':
                Dict['UE_Zeiten'] = '08:00 - 16:00'
            elif Dict[key] == 'Teilzeit':
                Dict['UE_Zeiten'] = '08:00 - 12:00'
        Dict['EmptyValue'] = NoInput
        return Dict
    
    # --- Copy the templates to a new directory in whicht the new files are saved ---
    def copy_templates(Dict) -> None:
        input = Dict['Vorlagen']
        output = Dict['Teilnehmerordner']
        shutil.copytree(input, output, dirs_exist_ok=True)
    
    # --- Create the new files and render them with the inputs from the dictionary ---
    def Create_files(Dict) -> None:
        for file in glob.iglob(Dict['Teilnehmerordner'] + f'{os.sep}**{os.sep}*.docx', recursive = True):
            doc = DocxTemplate(file)
            doc.render(Dict)
            doc.save(file)
    
    # --- Rename files with missing information '00_...' for ease of finding ---
    def Rename_Files(Dict) -> None:
        find = Dict["EmptyValue"]
        for file in glob.iglob(Dict['Teilnehmerordner'] + f'{os.sep}**{os.sep}*.docx', recursive = True):
            doc = Document(file)

            if [True for tbl in doc.tables for row in tbl.rows for cell in row.cells for p in cell.paragraphs if find in p.text]:
                os.rename(file,f'{os.path.dirname(file)}{os.sep}00_{Path(file).name}')

            if os.path.exists(file):
                for p in doc.paragraphs:
                    if find in p.text:
                        os.rename(file,f'{os.path.dirname(file)}{os.sep}00_{Path(file).name}')
                        break

class Paths(ctk.CTkFrame):
     def __init__(self, Master, **kwargs):
        super().__init__(Master, **kwargs)

        # --- Set the background to transparent as the main app has a given background color ---
        self.configure(fg_color='transparent')

        # --- Create input option the templatesfolder ---
        ctk.CTkLabel(self, text = "Vorlagen", fg_color="transparent", text_color=TextColor).grid(row = 0,column = 0, padx = (40,10), pady = (20,10), sticky = 'w')
        Templates = ctk.CTkEntry(self, corner_radius = 10, width = 400, fg_color="transparent", text_color=TextColor)
        Templates.grid(row = 0, column = 1, padx = (0,10), pady = (20,10), sticky = 'w')
        Templates.DictKey = 'Vorlagen'

        # --- Searchbutton to look through the directory directly rather than typing - the chosen folderpath will be input into the entrybox ---
        BtnInput = ctk.CTkButton(self, width = 10, corner_radius = 10, text = "...", fg_color = ButtonColor, hover_color = HoverColor, command = lambda: search('Vorlagen'))
        BtnInput.grid(row = 0, column = 2, pady = (20,10))

        # --- Create input option the outputsfolder ---
        ctk.CTkLabel(self, text = "Teilnehmerordner", fg_color="transparent", text_color=TextColor).grid(row = 0, column = 3, padx = (20,10), pady = (20,10), sticky = 'w')
        OutputFolder = ctk.CTkEntry(self, corner_radius=10,width=400, fg_color="transparent", text_color=TextColor)
        OutputFolder.grid(row = 0, column = 4, padx = (0,10), pady = (20,10), sticky = 'w')
        OutputFolder.DictKey = 'Teilnehmerordner'

        # --- Searchbutton to look through the directory directly rather than typing - the chosen folderpath will be input into the entrybox ---
        BtnOutput = ctk.CTkButton(self, width = 10, corner_radius = 10, text = "...", fg_color = ButtonColor, hover_color = HoverColor, command = lambda: search('Teilnehmerordner'))
        BtnOutput.grid(row = 0, column = 5, pady = (20,10))
        
        # --- Checkbox if a new directory should be created ---
        NewDir = ctk.CTkCheckBox(self, text = "Teilnehmerordner anlegen? (Mustermann, Max)", border_color=ButtonColor, hover_color = HoverColor, text_color=TextColor)
        NewDir.grid(row = 2, column = 0, padx = 40, pady = (0,10), columnspan = 2, sticky = 'w')
        NewDir.DictKey = "NeueOrdner"

        # --- Searchfunktion to replace given text in entrybox before placing the chosen folderpath ---
        def search(button) -> None:

            # --- check which button was pressed and act acordingly ---
            if button == 'Teilnehmerordner':
                OutputFolder.delete(0,ctk.END)
                OutputFolder.insert(0,askdirectory(title='Bitte Teilnehmerordner auswählen'))
                OutputFolder.xview(ctk.END)
            else:
                Templates.delete(0,ctk.END)
                Templates.insert(0,askdirectory(title='Bitte Vorlagenordner auswählen'))
                Templates.xview(ctk.END)

class PersonalInfo(ctk.CTkFrame):
        def __init__(self, Master, **kwargs):
            super().__init__(Master, **kwargs)

            # --- Set the background to transparent as the main app has a given background color ---
            self.configure(fg_color='transparent')

            # --- Create a frame with entries for all needed information about the given person ---

            # --- Create inputoption for fulltime or parttime ---
            ctk.CTkLabel(self, text="Anrede", fg_color="transparent", text_color=TextColor).grid(row=0, column=0, padx=(40,5), pady=(15,0), sticky='w')
            Anrede = ctk.CTkComboBox(self, values=['Herr', 'Frau', 'Sonstiges'], width=200, state= 'readonly', fg_color=BackgroundColor, text_color=TextColor)
            Anrede.grid(row=0, column=1,pady=(15,0), padx=(0,20))
            Anrede.set('Anrede')
            Anrede.DictKey = 'Anrede'

            # --- Create inputoption for name ---
            ctk.CTkLabel(self, text="Vorname", fg_color="transparent", text_color=TextColor).grid(row=1, column=0, padx=(40,5), pady=(15,0), sticky='w')
            Name = ctk.CTkEntry(self, corner_radius=10, width=200, fg_color="transparent", text_color=TextColor)
            Name.grid(row=1, column=1,pady=(15,0), padx=(0,20))
            Name.DictKey = 'Vorname'
            
            # --- Create inputoption for surname ---
            ctk.CTkLabel(self, text="Name", fg_color="transparent", text_color=TextColor).grid(row=2, column=0, padx=(40,5), pady=(15,0), sticky='w')
            Surname = ctk.CTkEntry(self, corner_radius=10, width=200, fg_color="transparent", text_color=TextColor)
            Surname.grid(row=2, column=1,pady=(15,0), padx=(0,20))
            Surname.DictKey = 'Name'

            # --- Create inputoption for phonenumber (homephone) ---
            ctk.CTkLabel(self, text="Geb. Datum", fg_color="transparent", text_color=TextColor).grid(row=3, column=0, padx=(40,5), pady=(15,0), sticky='w')
            Phone = ctk.CTkEntry(self, corner_radius=10, width=200, fg_color="transparent", text_color=TextColor)
            Phone.grid(row=3, column=1,pady=(15,0), padx=(0,20))
            Phone.DictKey = 'Geburtsdatum'

            # --- Create inputoption for phonenumber (homephone) ---
            ctk.CTkLabel(self, text="Staatsang.", fg_color="transparent", text_color=TextColor).grid(row=4, column=0, padx=(40,5), pady=(15,0), sticky='w')
            Phone = ctk.CTkEntry(self, corner_radius=10, width=200, fg_color="transparent", text_color=TextColor)
            Phone.grid(row=4, column=1,pady=(15,0), padx=(0,20))
            Phone.DictKey = 'Staatsang'

            # --- Create inputoption for streetadress ---
            ctk.CTkLabel(self, text="Straße", fg_color="transparent", text_color=TextColor).grid(row=5, column=0, padx=(40,5), pady=(15,0), sticky='w')
            Streetadress = ctk.CTkEntry(self, corner_radius=10, width=200, fg_color="transparent", text_color=TextColor)
            Streetadress.grid(row=5, column=1,pady=(15,0), padx=(0,20))
            Streetadress.DictKey = 'Straße'

            # --- Create inputoption for zipcode ---
            ctk.CTkLabel(self, text="Postleitzahl", fg_color="transparent", text_color=TextColor).grid(row=6, column=0, padx=(40,5), pady=(15,0), sticky='w')
            Zipcode = ctk.CTkEntry(self, corner_radius=10, width=200, fg_color="transparent", text_color=TextColor)
            Zipcode.grid(row=6, column=1,pady=(15,0), padx=(0,20))
            Zipcode.DictKey = 'PLZ'

            # --- Create inputoption for city ---
            ctk.CTkLabel(self, text="Ort", fg_color="transparent", text_color=TextColor).grid(row=7, column=0, padx=(40,5), pady=(15,0), sticky='w')
            City = ctk.CTkEntry(self, corner_radius=10, width=200, fg_color="transparent", text_color=TextColor)
            City.grid(row=7, column=1,pady=(15,0), padx=(0,20))
            City.DictKey = 'Ort'

            # --- Create inputoption for email ---
            ctk.CTkLabel(self, text="E-Mail", fg_color="transparent", text_color=TextColor).grid(row=8, column=0, padx=(40,5), pady=(15,0), sticky='w')
            Email = ctk.CTkEntry(self, corner_radius=10, width=200, fg_color="transparent", text_color=TextColor)
            Email.grid(row=8, column=1,pady=(15,0), padx=(0,20))
            Email.DictKey = 'EMail'

            # --- Create inputoption for mobilenumber ---
            ctk.CTkLabel(self, text="Mobil", fg_color="transparent", text_color=TextColor).grid(row=9, column=0, padx=(40,5), pady=(15,0), sticky='w')
            Mobile = ctk.CTkEntry(self, corner_radius=10, width=200, fg_color="transparent", text_color=TextColor)
            Mobile.grid(row=9, column=1,pady=(15,0), padx=(0,20))
            Mobile.DictKey = 'Mobil'

            # --- Create inputoption for phonenumber (homephone) ---
            ctk.CTkLabel(self, text="Telefon", fg_color="transparent", text_color=TextColor).grid(row=10, column=0, padx=(40,5), pady=(15,0), sticky='w')
            Phone = ctk.CTkEntry(self, corner_radius=10, width=200, fg_color="transparent", text_color=TextColor)
            Phone.grid(row=10, column=1,pady=(15,0), padx=(0,20))
            Phone.DictKey = 'Telefon'

class CourseInfo(ctk.CTkFrame):
        def __init__(self, Master, **kwargs):
            super().__init__(Master, **kwargs)

            # --- Set the background to transparent as the main app has a given background color ---
            self.configure(fg_color='transparent')
            
            # --- Create a frame with entries for all needed information about the given course ---

            # --- Create inputoption for coursename ---
            ctk.CTkLabel(self, text="Kundennummer", fg_color="transparent", text_color=TextColor).grid(row=0, column=0, padx=(20,5), pady=(15,0), sticky='w')
            CourseName = ctk.CTkEntry(self, corner_radius=10, width=200, fg_color="transparent", text_color=TextColor)
            CourseName.grid(row=0, column=1,pady=(15,0), padx=(0,20))
            CourseName.DictKey = 'Kundennummer'
            
            # --- Create inputoption for coursename ---
            ctk.CTkLabel(self, text="Kursnamen", fg_color="transparent", text_color=TextColor).grid(row=1, column=0, padx=(20,5), pady=(15,0), sticky='w')
            CourseName = ctk.CTkEntry(self, corner_radius=10, width=200, fg_color="transparent", text_color=TextColor)
            CourseName.grid(row=1, column=1,pady=(15,0), padx=(0,20))
            CourseName.DictKey = 'Kursnamen'
            
            # --- Create inputoption for coursenumber ---
            ctk.CTkLabel(self, text="KursID", fg_color="transparent", text_color=TextColor).grid(row=2, column=0, padx=(20,5), pady=(15,0), sticky='w')
            CourseNumber = ctk.CTkEntry(self, corner_radius=10, width=200, fg_color="transparent", text_color=TextColor)
            CourseNumber.grid(row=2, column=1,pady=(15,0), padx=(0,20))
            CourseNumber.DictKey = 'KursID'

            # --- Create inputoption for the course startdate ---
            ctk.CTkLabel(self, text="Eintritt", fg_color="transparent", text_color=TextColor).grid(row=3, column=0, padx=(20,5), pady=(15,0), sticky='w')
            StartDate = ctk.CTkEntry(self, corner_radius=10, width=200, fg_color="transparent", text_color=TextColor)
            StartDate.grid(row=3, column=1,pady=(15,0), padx=(0,20))
            StartDate.DictKey = 'Eintritt'

            # --- Create inputoption for the course enddate ---
            ctk.CTkLabel(self, text="Austritt", fg_color="transparent", text_color=TextColor).grid(row=4, column=0, padx=(20,5), pady=(15,0), sticky='w')
            EndDate = ctk.CTkEntry(self, corner_radius=10, width=200, fg_color="transparent", text_color=TextColor)
            EndDate.grid(row=4, column=1,pady=(15,0), padx=(0,20))
            EndDate.DictKey = 'Austritt'

            # --- Create inputoption for fulltime or parttime ---
            ctk.CTkLabel(self, text="Voll-/Teilzeit", fg_color="transparent", text_color=TextColor).grid(row=5, column=0, padx=(20,5), pady=(15,0), sticky='w')
            Fulltime = ctk.CTkComboBox(self, values=['Vollzeit', 'Teilzeit'], width=200, state= 'readonly', fg_color=BackgroundColor, text_color=TextColor)
            Fulltime.grid(row=5, column=1,pady=(15,0), padx=(0,20))
            Fulltime.set('Vollzeit')
            Fulltime.DictKey = 'VZ_TZ'

            # --- Create inputoption for course hours ---
            ctk.CTkLabel(self, text="UE's", fg_color="transparent", text_color=TextColor).grid(row=6, column=0, padx=(20,5), pady=(15,0), sticky='w')
            Hours = ctk.CTkEntry(self, corner_radius=10, width=200, fg_color="transparent", text_color=TextColor)
            Hours.grid(row=6, column=1,pady=(15,0), padx=(0,20))
            Hours.DictKey = 'UE'

            # --- Create inputoption for fulltime or parttime ---
            ctk.CTkLabel(self, text="Praktikum", fg_color="transparent", text_color=TextColor).grid(row=7, column=0, padx=(20,5), pady=(15,0), sticky='w')
            Fulltime = ctk.CTkComboBox(self, values=['6 Monate', '9 Monate', 'Kein Praktikum'], width=200, state= 'readonly', fg_color=BackgroundColor, text_color=TextColor)
            Fulltime.grid(row=7, column=1,pady=(15,0), padx=(0,20))
            Fulltime.set('6 Monate')
            Fulltime.DictKey = 'Praktikum'

            # --- Create inputoption for course hours ---
            ctk.CTkLabel(self, text="Kostenträger", fg_color="transparent", text_color=TextColor).grid(row=8, column=0, padx=(20,5), pady=(15,0), sticky='w')
            Hours = ctk.CTkEntry(self, corner_radius=10, width=200, fg_color="transparent", text_color=TextColor)
            Hours.grid(row=8, column=1,pady=(15,0), padx=(0,20))
            Hours.DictKey = 'Kostentraeger'

            # --- Create inputoption for course hours ---
            ctk.CTkLabel(self, text="Kursgebühren", fg_color="transparent", text_color=TextColor).grid(row=9, column=0, padx=(20,5), pady=(15,0), sticky='w')
            Hours = ctk.CTkEntry(self, corner_radius=10, width=200, fg_color="transparent", text_color=TextColor)
            Hours.grid(row=9, column=1,pady=(15,0), padx=(0,20))
            Hours.DictKey = 'Kursgebuehren'

            # --- Create inputoption for course hours ---
            ctk.CTkLabel(self, text="Vermittler*in", fg_color="transparent", text_color=TextColor).grid(row=10, column=0, padx=(20,5), pady=(15,0), sticky='w')
            Hours = ctk.CTkEntry(self, corner_radius=10, width=200, fg_color="transparent", text_color=TextColor)
            Hours.grid(row=10, column=1,pady=(15,0), padx=(0,20))
            Hours.DictKey = 'Vermittlerin'

class MyApp(ctk.CTk):
        def __init__(self):
            super().__init__()

            # --- Create appwindow and set the parameters needed ---
            self.geometry("1150x680")
            self.resizable(width=False, height=False)
            self.title("Viona")
            self.configure(fg_color=BackgroundColor)
            self.iconbitmap("logo.ico")


            # --- Add infos ---

            # --- Pathinformation ---
            paths = Paths(Master=self)
            paths.grid(row=0, column=0, columnspan = 2, sticky="nsew")

            # --- Personal infos ---
            ctk.CTkLabel(self, text='Personal Infos: ', font=("none", 20), text_color=TextColor).grid(row=1, column=0, pady=(0,10))
            personinfo = PersonalInfo(Master=self)
            personinfo.grid(row=2, column=0, pady=(0,20), sticky="nsew")

            # --- Course infos ---
            ctk.CTkLabel(self, text='Course Infos: ', font=("none", 20), text_color=TextColor).grid(row=1, column=1, pady=(0,10))
            courseinfo = CourseInfo(Master=self)
            courseinfo.grid(row=2,column=1, padx=40, pady=(0,20), sticky="nsew")

            # --- Buttons (Cancel and Submit) ---
            ctk.CTkButton(self,text="Cancel",fg_color = ButtonColor, hover_color = HoverColor, width=200, command=self.destroy).grid(row=3,column=0, padx= (120,0), sticky='w')
            ctk.CTkButton(self,text="Submit",fg_color = ButtonColor, hover_color = HoverColor, width=200, command =self.submit).grid(row=3,column=1, padx= (150,0), sticky='w')


        # --- Check all entries and save the values to a dictionary
        def submit(appwindow) -> None:
            dict = {}
            try:
                # --- due to the class structure we loop through the frames of the appwindow ---
                for child in appwindow.winfo_children():

                    # --- Loop through the widgets of each frame on the appwindow ---              
                    for widget in child.winfo_children():

                        # --- If the Widget is an Entry, a Combobox or a Checkbox we save the value in a dictionary ---
                        # --- DictKey is the "name" of each widget - used to seperate the values in the dictionary later on ---
                        if isinstance(widget, ctk.CTkEntry) or isinstance(widget, ctk.CTkComboBox) or isinstance(widget, ctk.CTkCheckBox):
                            dict[widget.DictKey] = widget.get()

                # --- If any values are empty, show popup with info and provide solution for unknown values ---
                if any(i == "" for i in dict.values()):
                    Popup(Title = 'Missing information',
                        Message = 'Bitte die noch fehlende Informationen ergänzen. Liegt Ihnen die Information nicht vor machen Sie bitte ein Minus" - " anstatt!',
                        Icon = "warning")
                    
                # --- If all information is provided ---
                else: 

                    # --- Check if a new directory is needed and update the dictionary if necessary ---
                    dict = HandleSubmit.New_dir(dict)

                    # --- If all inputs have been provided and the given paths for the templates and the output exists continue ---
                    # --- Otherwise an errormessage is show ---
                    if HandleSubmit.Test_path(dict) == True:

                        # --- Cleanup the dictionary and replace missing values for later formatting ---
                        Context = HandleSubmit.Clean_dict(dict)

                        # --- Copy templatedirectory to new directory ---
                        HandleSubmit.copy_templates(Context)

                        # --- Render files from the given templates ---
                        HandleSubmit.Create_files(Context)

                        # --- Rename files with missing information '00_...' for ease of finding ---
                        HandleSubmit.Rename_Files(Context)

                        # --- Show Succesmessage ---
                        if Succes() == True:
                            appwindow.quit()

            except Exception as error:

                Popup(
                    Title='Problem',
                    Message=f'Es ist ein Problem aufgetreten. \n Folgende Fehler ist aufgetreten: {error}',
                    Icon='warning'
                )
            
# --- Popupwindow to show messages ---
def Popup(Title, Message, Icon, button='OK') -> str:
    popup = CTkMessagebox(width=500,
                  height=200,
                  title=Title,
                  message=Message, 
                  icon=Icon,
                  option_1=button,
                  font=(None,15),
                  border_width=2,
                  border_color=TextColor,
                  bg_color=BackgroundColor,
                  fg_color=BackgroundColor,
                  text_color=TextColor,
                  title_color=TextColor)
    return popup.get()

def Succes() -> bool:
    quit = Popup(Title = 'Succes',
            Message = 'All your documents have been created!',
            Icon = "check")
    if quit == 'OK':
        return True

def main():
    app = MyApp()
    app.mainloop()

if __name__ == "__main__":
    main()